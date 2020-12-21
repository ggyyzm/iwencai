from PySide2.QtUiTools import QUiLoader
from PySide2.QtWidgets import QApplication, QMessageBox, QCheckBox, QTableWidgetItem, QTextBrowser
from PySide2.QtCore import Signal, QObject
from PySide2.QtGui import QPixmap
from grabInfov1 import GrabStockInfo
from docx import Document
from docx.shared import Pt, Cm
import os
import time
import threading
# import win32api
# import win32print
import base64

# dir_name = os.path.dirname(PySide2.__file__)
# plugin_path = os.path.join(dir_name, 'plugins', 'platforms')
# os.environ['QT_QPA_PLATFORM_PLUGIN_PATH'] = plugin_path
# dirname = os.path.dirname(__file__)
# plugin_path = os.path.join(dirname, 'plugins', 'platforms')
# os.environ['QT_QPA_PLATFORM_PLUGIN_PATH'] = plugin_path


# 自定义信号源对象类型，一定要继承自 QObject
class MySignals(QObject):
    # 定义一种信号，两个参数 类型分别是： QTextBrowser 和 字符串
    text_print = Signal(QTextBrowser, str)
    next_page = Signal()
    select_all = Signal()
    prevent_continuous = Signal()


class mainUI:
    def __init__(self):
        # 加载界面
        self.ui = QUiLoader().load('ui/main.ui')

        # 属性
        self.selected_num = 0       # 已选股票数
        self.page_searched_num = 0  # 当前页搜索到的股票数
        self.searched_num = 0       # 搜索到的股票数
        self.isSelectedAll = False  # 是否全选
        self.current_page = 1   # 当前页
        self.all_pages = 1      # 所有页
        self.table0_rows = 7    # 第一个表格行数
        self.table0_cols = 2
        self.table1_rows = 5    # 第二个表格行数
        self.table1_cols = 4
        self.table2_rows = 2    # 第三个表格行数
        self.table2_cols = 1
        self.docx_save_dir = "./文档存储"
        self.img_save_dir = "./image_cache"
        self.grab = None
        self.today_date = ""
        self.lock = threading.Lock()
        self.thread = None
        self.is_login = False
        # self.event = []

        # 实例化
        self.ms = MySignals()
        # 自定义信号的处理函数
        self.ms.text_print.connect(self.printToGui)
        self.ms.next_page.connect(self.next_page_button_function)
        self.ms.select_all.connect(self.select_all_button_function)
        self.ms.prevent_continuous.connect(self.prevent_continuous_clicks)

        # 控件事件
        self.ui.iwencai_button.clicked.connect(self.iwencai_button_function)
        self.ui.select_all_button.clicked.connect(self.select_all_button_function)
        self.ui.print_button.clicked.connect(self.print_button_function)
        self.ui.previous_page_button.clicked.connect(self.previous_page_button_function)
        self.ui.next_page_button.clicked.connect(self.next_page_button_function)
        self.ui.thelast_page_button.clicked.connect(self.thelast_page_button_function)
        self.ui.print_all_button.clicked.connect(self.print_all_button_function)

        # 子窗口
        self.sub_ui = QUiLoader().load('ui/login.ui')
        self.sub_ui.buttonBox.accepted.connect(self.sub_ui_login)

    # 子窗口点击OK事件
    def sub_ui_login(self):
        self.is_login, message = self.grab.do_login(self.sub_ui.username_edit.text(), self.sub_ui.passwd_edit.text(),
                           self.sub_ui.veri_img_edit.text())
        if self.is_login is False:
            self.print_show(self.ui.print_show, "登录失败：" + message)
            self.get_login_info()

    # 点击上一页按钮事件
    def previous_page_button_function(self):
        stocklist = self.grab.getStocksListByPage(self.current_page - 1)
        if stocklist is None:
            # QMessageBox.warning(self.ui, "搜索不到内容", "搜索不到查找内容，请重新搜索！")
            self.print_show(self.ui.print_show, "搜索不到查找内容，请重新搜索！")
            return
        self.update_main_UI(stocklist)

    # 点击下一页按钮事件
    def next_page_button_function(self):
        stocklist = self.grab.getStocksListByPage(self.current_page + 1)
        print('next_page_button_function: self.current_page = ' + str(stocklist['current_page']))
        if stocklist is None:
            self.print_show(self.ui.print_show, "搜索不到查找内容，请重新搜索！")
            return
        self.update_main_UI(stocklist)

    # 点击最后一页按钮事件
    def thelast_page_button_function(self):
        stocklist = self.grab.getStocksListByPage(self.all_pages)
        if stocklist is None:
            self.print_show(self.ui.print_show, "搜索不到查找内容，请重新搜索！")
            return
        self.update_main_UI(stocklist)

    # 点击i问财搜索按钮事件
    def iwencai_button_function(self):
        # 关闭driver
        try:
            self.grab.driver.close()
        except:
            pass

        search_content = self.ui.search_input.text()
        self.grab = GrabStockInfo(keywords=[search_content], driverPath=r'.\chromedriver.exe', showWindow=True)
        time.sleep(1)
        # 获取当前页面的列表
        stocklist = self.grab.allStockList()
        if stocklist is None:
            QMessageBox.warning(self.ui, "搜索不到内容", "搜索不到查找内容，请重新搜索！")
            return

        # 更新主界面信息
        self.update_main_UI(stocklist)

        # 登录
        self.is_login = self.grab.load_cookies()
        time.sleep(0.5)
        if self.is_login is True:
            pass
        else:
            self.grab.login()
            self.get_login_info()

    # 更新主界面
    def update_main_UI(self, stocklist):
        # 先清空
        self.ui.searched_show.setRowCount(0)
        self.ui.searched_show.clear()
        self.selected_num = 0
        self.ui.selected_num.setText(str(self.selected_num))
        self.ui.current_page.clear()
        self.ui.all_page_num.clear()
        self.isSelectedAll = False

        # 显示搜索到的总条目数/当前页条目数
        self.searched_num = int(stocklist['all_stocks'])
        self.ui.searched_num.setText(str(self.searched_num))
        self.page_searched_num = int(stocklist['name_and_code'].__len__())
        self.ui.page_searched_num.setText(str(self.page_searched_num))

        # 在显示框中显示搜索到的股票信息
        for i in range(stocklist['name_and_code'].__len__()):
            self.ui.searched_show.insertRow(i)
            checkBox = QCheckBox()
            checkBox.toggled.connect(lambda isChecked: self.check_box_checked(isChecked))
            textItem_0 = QTableWidgetItem(stocklist['name_and_code'][i][0])
            textItem_1 = QTableWidgetItem(stocklist['name_and_code'][i][1])

            self.ui.searched_show.setCellWidget(i, 0, checkBox)
            self.ui.searched_show.setItem(i, 1, textItem_0)
            self.ui.searched_show.setItem(i, 2, textItem_1)

        # 更新按钮enabled
        try:
            self.current_page = int(stocklist['current_page'])
            self.all_pages = int(stocklist['all_pages'])
        except:
            pass
        self.update_page_button()

    # 根据当前页面与总页面数更新按钮enabled
    def update_page_button(self):
        if self.current_page == self.all_pages:
            self.ui.next_page_button.setEnabled(False)
            self.ui.thelast_page_button.setEnabled(False)
        elif self.current_page != self.all_pages:
            self.ui.next_page_button.setEnabled(True)
            self.ui.thelast_page_button.setEnabled(True)
        else:
            self.print_show(self.ui.print_show, "错误：" + "The comparison between current_page and all_pages error!")
            # raise AttributeError("The comparison between current_page and all_pages error!")

        if self.current_page == 1:
            self.ui.previous_page_button.setEnabled(False)
        elif self.current_page > 1:
            self.ui.previous_page_button.setEnabled(True)
        else:
            self.print_show(self.ui.print_show, "错误：" + "The comparison between current_page and 1 error!")
            # raise AttributeError("The comparison between current_page and 1 error!")

        self.ui.current_page.setText(str(self.current_page))
        self.ui.all_page_num.setText(str(self.all_pages))

    # 弹出登录界面，获取登录信息
    def get_login_info(self):
        self.sub_ui.veri_img_label.setPixmap(QPixmap('veri.png'))
        self.sub_ui.show()

    # 选中/取消选中股票事件
    def check_box_checked(self, isChecked):
        if isChecked is True:
            self.selected_num += 1
            self.ui.selected_num.setText(str(self.selected_num))
            if self.selected_num == self.page_searched_num:
                self.isSelectedAll = True
        else:
            self.selected_num -= 1
            self.ui.selected_num.setText(str(self.selected_num))
            if self.selected_num < self.page_searched_num:
                self.isSelectedAll = False
            else:
                self.print_show(self.ui.print_show, "错误：" + "selected_num & page_searched_num error")
                # raise AttributeError('selected_num & page_searched_num error')

    # 点击全选/全不选按钮事件
    def select_all_button_function(self):
        if self.isSelectedAll is False:
            for i in range(self.page_searched_num):
                self.ui.searched_show.cellWidget(i, 0).setChecked(True)
            # self.ui.searched_show.cellWidget(1, 0).setChecked(True)
            self.isSelectedAll = True
        else:
            for i in range(self.page_searched_num):
                self.ui.searched_show.cellWidget(i, 0).setChecked(False)
            self.isSelectedAll = False

    # 点击存档按钮事件
    def print_button_function(self):
        self.ui.iwencai_button.setEnabled(False)
        self.ui.print_button.setEnabled(False)
        self.ui.print_all_button.setEnabled(False)

        # 判断选中内容不可为空
        if self.selected_num == 0:
            self.print_show(self.ui.print_show, "没有选中需要存档的股票信息！")
            return
        else:
            pass

        # # 防止连点
        # self.prevent_continuous_thread()

        # 在项目目录下创建文档存储文件夹
        if os.path.exists(self.docx_save_dir):
            pass
        else:
            os.mkdir(self.docx_save_dir)

        # 在文档存储目录下创建当前日期文件夹
        self.today_date = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(time.time())).split()[0]
        if os.path.exists(self.docx_save_dir + '/' + self.today_date):
            pass
        else:
            os.mkdir(self.docx_save_dir + '/' + self.today_date)

        # 在项目目录下创建image文件夹
        if os.path.exists(self.img_save_dir):
            pass
        else:
            os.mkdir(self.img_save_dir)

        selected_list = []
        for i in range(self.page_searched_num):
            if self.ui.searched_show.cellWidget(i, 0).isChecked() is True:
                selected_list.append(True)
            else:
                selected_list.append(False)

        self.thread = threading.Thread(target=mainUI.Thread_func, args=(self, selected_list))
        self.thread.start()

    # 创建获取已选的股票的股票信息的线程
    def Thread_func(self, _list):
        # self.lock.acquire()
        for data in self.grab.getRetrivelStockInfo(_list):
            if len(data) == 2:
                self.save_multiple_info(data)
            elif len(data) == 14:
                self.save_single_info(data)
            else:
                QMessageBox.critical(self.ui, "错误", "data.len is error")
                raise AttributeError('data.len is error')
        # self.lock.release()
        self.prevent_continuous_thread()

    # 点击全部存档按钮事件
    def print_all_button_function(self):
        self.ui.iwencai_button.setEnabled(False)
        self.ui.print_button.setEnabled(False)
        self.ui.print_all_button.setEnabled(False)

        # 需要登录
        if self.is_login is False:
            QMessageBox.warning(self.ui, "错误", "请登录")
            return
        else:
            pass

        # 需要在第一页点此按钮
        if self.current_page != 1:
            QMessageBox.warning(self.ui, "错误", "请跳到第一页点击此按钮")
            return
        else:
            pass

        # 防止连点
        # self.prevent_continuous_thread()

        # 先全选
        if self.isSelectedAll is False:
            for i in range(self.page_searched_num):
                self.ui.searched_show.cellWidget(i, 0).setChecked(True)
            self.selected_num = self.page_searched_num
        else:
            pass
        # self.ui.searched_show.cellWidget(1, 0).setChecked(True)
        # self.selected_num = self.page_searched_num
        time.sleep(1)

        for i in range(self.all_pages):
            page_thread = threading.Thread(target=mainUI.page_Thread_func, args=(self, ))
            page_thread.start()

    # 翻页保存thread
    def page_Thread_func(self):
        self.lock.acquire()
        current_page_cache = self.current_page
        all_page_cache = self.all_pages
        self.print_button_function()
        self.thread.join()
        print('self.current_page = ' + str(self.current_page))
        print('self.all_pages = ' + str(self.all_pages))
        if all_page_cache != self.all_pages:
            self.print_show(self.ui.print_show, "网页更新，总页数已改变，请手选存档，或重新全部存档")

        if self.current_page != self.all_pages:
            self.next_page_thread()
            time.sleep(10)
            if self.current_page == 1:
                self.print_show(self.ui.print_show, "网页更新，正在跳转，请稍等")
                while (current_page_cache + 1) != self.current_page:
                    self.next_page_thread()
                    time.sleep(10)
            else:
                pass
        else:
            pass

        print('self.current_page = ' + str(self.current_page))
        print('self.all_pages = ' + str(self.all_pages))
        if current_page_cache == self.all_pages:
            pass
        else:
            self.select_all_thread()
            time.sleep(1)
        self.prevent_continuous_thread()
        self.lock.release()

    # 在子线程中翻下一页
    def next_page_thread(self):
        def _next_page_thread():
            # 通过Signal 的 emit 触发执行 主线程里面的处理函数
            # emit参数和定义Signal的数量、类型必须一致
            self.ms.next_page.emit()
        threading.Thread(target=_next_page_thread).start()

    # 在子线程中全选
    def select_all_thread(self):
        def _select_all_thread():
            # 通过Signal 的 emit 触发执行 主线程里面的处理函数
            # emit参数和定义Signal的数量、类型必须一致
            self.ms.select_all.emit()
        threading.Thread(target=_select_all_thread).start()

    # 用子线程防止连点
    def prevent_continuous_thread(self):
        def _prevent_continuous_thread():
            # 通过Signal 的 emit 触发执行 主线程里面的处理函数
            # emit参数和定义Signal的数量、类型必须一致
            self.ms.prevent_continuous.emit()
        threading.Thread(target=_prevent_continuous_thread).start()

    # 防止连点
    def prevent_continuous_clicks(self):
        self.ui.iwencai_button.setEnabled(False)
        self.ui.print_button.setEnabled(False)
        self.ui.print_all_button.setEnabled(False)
        time.sleep(2)
        self.ui.iwencai_button.setEnabled(True)
        self.ui.print_button.setEnabled(True)
        self.ui.print_all_button.setEnabled(True)

    # 保存搜索单个股票时的data到word(无图)
    def save_single_info(self, data):
        file = Document()

        table_0 = file.add_table(rows=self.table0_rows, cols=self.table0_cols, style='Table Grid')
        table_1 = file.add_table(rows=self.table1_rows, cols=self.table1_cols, style='Table Grid')
        table_2 = file.add_table(rows=self.table2_rows, cols=self.table2_cols, style='Table Grid')

        table_0.style.font.size = Pt(14)
        table_1.style.font.size = Pt(14)
        table_2.style.font.size = Pt(14)

        # table_0
        table_0.cell(0, 0).text = "股票简称：" + data['stock_abbreviation']
        table_0.cell(0, 1).text = "股票代码：" + data['stock_code']
        table_0.cell(1, 0).text = "涨跌幅：" + data['pick_price'] + '  ' + data['price_limit']
        try:
            table_0.cell(1, 1).text = data['reasons_limit'][0][2].split('，')[1]
        except:
            table_0.cell(1, 1).text = "缺少涨跌停原因"
        table_0.cell(2, 0).text = "总股本(股)：" + data['general_capital']
        table_0.cell(2, 1).text = "流通比例(%)：" + data['circulation_proportion']
        table_0.cell(3, 0).text = "总市值(亿元)：" + data['total_market_value']
        table_0.cell(3, 1).text = "流通市值(亿元)：" + data['market_value']
        table_0.cell(4, 0).text = "市盈率(倍)：" + data['pe_ratio']
        table_0.cell(4, 1).text = "市净率(倍)：" + data['net_ratio']
        # table_0的四张图
        table_0.cell(5, 0).text = "分时图"
        table_0.cell(5, 1).text = "日K图"
        table_0.cell(6, 0).text = "周K图"
        table_0.cell(6, 1).text = "月K图"

        # table_1
        table_1.cell(0, 0).text = "报告期时间"
        table_1.cell(0, 1).text = "净利润(元)"
        table_1.cell(0, 2).text = "净利润同比(%)"
        table_1.cell(0, 3).text = "每股收益(元)"
        for i in range(1, 5):
            for j in range(4):
                table_1.cell(i, j).text = data['financial_all'][i-1][j]

        # table_2
        string = data['main_product_name'][0]
        for i in range(1, len(data['main_product_name'])):
            string += "，" + data['main_product_name'][i]
        table_2.cell(0, 0).text = "主营产品名称：" + string
        try:
            string = data['affiliation_concept'][0]
            for i in range(1, len(data['affiliation_concept'])):
                string += "，" + data['affiliation_concept'][i]
            table_2.cell(1, 0).text = "所属概念：" + string
        except:
            table_2.cell(1, 0).text = "缺少所属概念"

        # word文件名不能带'*'
        if data["stock_abbreviation"][0] == '*':
            file_name = '星号+' + data["stock_abbreviation"].split('*')[1]
        else:
            file_name = data["stock_abbreviation"]

        try:
            file.save(self.docx_save_dir + "/" + file_name + "_" + self.today_date + ".docx")
        except:
            QMessageBox.warning(self.ui, "无法保存文件", "无法保存" + file_name + "！请检查是否已打开同名文件但未关闭。")
            # return
        # 保存成功，输出到显示框中
        self.print_show(self.ui.print_show, data["stock_abbreviation"] + ".docx" + "已保存")

        # 打印
        # self.lock.acquire()
        # self.printer_loading(self.docx_save_dir + "/" + data["stock_code"] + ".docx")
        # 打印成功
        # self.print_show(self.ui.print_show, data["stock_abbreviation"] + ".docx" + "已打印")
        # self.lock.release()

    # 保存搜索多个股票时的data到word(有图)
    def save_multiple_info(self, data):
        file = Document()

        table_0 = file.add_table(rows=self.table0_rows, cols=self.table0_cols, style='Table Grid')
        table_1 = file.add_table(rows=self.table1_rows, cols=self.table1_cols, style='Table Grid')
        table_2 = file.add_table(rows=self.table2_rows, cols=self.table2_cols, style='Table Grid')

        table_0.style.font.size = Pt(11)
        table_1.style.font.size = Pt(11)
        table_2.style.font.size = Pt(11)

        # table_0
        table_0.cell(0, 0).text = "股票简称：" + data[0]['stock_abbreviation']
        table_0.cell(0, 1).text = "股票代码：" + data[0]['stock_code']
        table_0.cell(1, 0).text = "涨跌幅：" + data[0]['pick_price'] + '  ' + data[0]['price_limit']
        try:
            table_0.cell(1, 1).text = data[0]['reasons_limit'][0][2].split('，')[1]
        except:
            table_0.cell(1, 1).text = "缺少涨跌停原因"
        table_0.cell(2, 0).text = "总股本(股)：" + data[0]['general_capital']
        table_0.cell(2, 1).text = "流通比例(%)：" + data[0]['circulation_proportion']
        table_0.cell(3, 0).text = "总市值(亿元)：" + data[0]['total_market_value']
        table_0.cell(3, 1).text = "流通市值(亿元)：" + data[0]['market_value']
        table_0.cell(4, 0).text = "市盈率(倍)：" + data[0]['pe_ratio']
        table_0.cell(4, 1).text = "市净率(倍)：" + data[0]['net_ratio']

        # table_0的四张图
        for i in range(4):
            image = open(self.img_save_dir + "/" + str(i) + ".jpg", 'wb')
            image.write(base64.b64decode(data[1][i]))
            image.close()
        # table_0.cell(5, 0).text = "分时图"
        run = table_0.cell(5, 0).paragraphs[0].add_run()
        run.add_picture(self.img_save_dir + "/0.jpg", width=Cm(8), height=Cm(7))
        # table_0.cell(5, 1).text = "日K图"
        run = table_0.cell(5, 1).paragraphs[0].add_run()
        run.add_picture(self.img_save_dir + "/1.jpg", width=Cm(8), height=Cm(7))
        # table_0.cell(6, 0).text = "周K图"
        run = table_0.cell(6, 0).paragraphs[0].add_run()
        run.add_picture(self.img_save_dir + "/2.jpg", width=Cm(8), height=Cm(7))
        # table_0.cell(6, 1).text = "月K图"
        run = table_0.cell(6, 1).paragraphs[0].add_run()
        run.add_picture(self.img_save_dir + "/3.jpg", width=Cm(8), height=Cm(7))

        # table_1
        table_1.cell(0, 0).text = "报告期时间"
        table_1.cell(0, 1).text = "净利润(元)"
        table_1.cell(0, 2).text = "净利润同比(%)"
        table_1.cell(0, 3).text = "每股收益(元)"
        for i in range(1, 5):
            for j in range(4):
                table_1.cell(i, j).text = data[0]['financial_all'][i - 1][j]

        # table_2
        string = data[0]['main_product_name'][0]
        for i in range(1, len(data[0]['main_product_name'])):
            string += "，" + data[0]['main_product_name'][i]
        table_2.cell(0, 0).text = "主营产品名称：" + string
        try:
            string = data[0]['affiliation_concept'][0]
            for i in range(1, len(data[0]['affiliation_concept'])):
                string += "，" + data[0]['affiliation_concept'][i]
            table_2.cell(1, 0).text = "所属概念：" + string
        except:
            table_2.cell(1, 0).text = "缺少所属概念"

        # word文件名不能带'*'
        if data[0]["stock_abbreviation"][0] == '*':
            file_name = '星号+' + data[0]["stock_abbreviation"].split('*')[1]
        else:
            file_name = data[0]["stock_abbreviation"]

        # 保存.docx文件
        try:
            file.save(self.docx_save_dir + "/" + self.today_date + "/" + file_name + "_" + self.today_date + ".docx")
        except:
            self.print_show(self.ui.print_show, "无法保存" + file_name + "！请检查是否已打开同名文件但未关闭。")
        #     return
        # 保存成功，输出到显示框中
        self.print_show(self.ui.print_show, data[0]["stock_abbreviation"] + ".docx" + "已保存到" + self.docx_save_dir + '/' + self.today_date + "文件夹")

        # 打印
        # self.lock.acquire()
        # self.printer_loading(self.docx_save_dir + "/" + file_name + ".docx")
        # 打印成功
        # self.print_show(self.ui.print_show, data[0]["stock_abbreviation"] + ".docx" + "已打印")
        # self.lock.release()

    # # 打印
    # def printer_loading(self, filename):
    #     try:
    #         open_file = open(filename, "r")
    #     except:
    #         QMessageBox.warning(self.ui, "无法打开文件", "无法打开" + filename + "！请检查是否已打开同名文件但未关闭。")
    #         return
    #
    #     filename_path = os.path.join(os.getcwd() + '/' + filename.split('.', 1)[1])
    #     # open_file.close()
    #     # return  # 测试，不浪费纸
    #     win32api.ShellExecute(
    #         0,
    #         "print",
    #         filename_path,
    #         #
    #         # If this is None, the default printer will
    #         # be used anyway.
    #         #
    #         '/d:"%s"' % win32print.GetDefaultPrinter(),
    #         ".",
    #         0
    #     )
    #     open_file.close()

    # 输出信息到显示框
    def print_show(self, textBrowser, text):
        def _print_show_thread():
            # 通过Signal 的 emit 触发执行 主线程里面的处理函数
            # emit参数和定义Signal的数量、类型必须一致
            self.ms.text_print.emit(textBrowser, text)
        threading.Thread(target=_print_show_thread).start()

    def printToGui(self, textBrowser, text):
        textBrowser.append(text)  # 文本框逐条添加数据
        textBrowser.moveCursor(textBrowser.textCursor().End)  # 文本框显示到底部


app = QApplication([])
main_UI = mainUI()
main_UI.ui.show()
app.exec_()
